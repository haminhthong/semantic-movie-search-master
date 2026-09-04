from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("reports/huong_dan_cai_thien_du_an_4_tang.docx")
BLUE = "2E74B5"
DARK = "0B2545"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT = "E8EEF5"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"


def font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
        shade_cell(table.cell(0, i), LIGHT)
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    table_geometry(table, widths)
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def callout(doc, label, text, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p._p.get_or_add_pPr().append(shd)
    font(p.add_run(f"{label}: "), 10.5, color, True)
    font(p.add_run(text), 10.5, DARK)


def checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
for list_style in ("List Bullet", "List Bullet 2", "List Number"):
    style = doc.styles[list_style]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25
doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
doc.styles["List Number"].paragraph_format.left_indent = Inches(0.375)
doc.styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)

header = section.header.paragraphs[0]
font(header.add_run("MOVIESCOUT AI  |  HƯỚNG DẪN CẢI THIỆN 4 TẦNG"), 9, MUTED, True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("Bản hướng dẫn kỹ thuật • 01/09/2026"), 9, MUTED)

# Editorial cover
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("TECHNICAL IMPROVEMENT GUIDE"), 10.5, GOLD, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("Hướng dẫn cải thiện dự án theo 4 tầng"), 29, DARK, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Problem → AI/ML Correctness → Software Engineering → Production & Business Value"), 14, DARK_BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(26)
font(p.add_run("Semantic Movie Search / MovieScout AI"), 12, DARK, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Tài liệu hành động từ audit repository hiện tại"), 10, MUTED, False, True)
doc.add_page_break()

doc.add_heading("Mục lục", level=1)
for item in [
    "1. Phạm vi và kết luận audit",
    "2. Tầng 1 — Problem",
    "3. Tầng 2 — AI/ML correctness",
    "4. Tầng 3 — Software Engineering",
    "5. Tầng 4 — Production / Business value",
    "6. Lộ trình triển khai theo phase",
    "7. Definition of Done và checklist nghiệm thu",
]:
    bullet(doc, item)

doc.add_heading("1. Phạm vi và kết luận audit", level=1)
doc.add_paragraph(
    "Tài liệu này chuyển kết quả audit repository thành kế hoạch cải thiện có thể thực hiện. Mục tiêu không phải "
    "thêm nhiều thành phần AI hơn, mà là chứng minh hệ thống giải quyết đúng bài toán, đo lường đáng tin, chạy lại "
    "được trên máy khác và có đường nâng cấp rõ ràng lên môi trường production."
)
add_table(
    doc,
    ["Tầng", "Hiện trạng", "Mục tiêu"],
    [
        ("Problem", "Đạt", "Khóa use case và tiêu chí thành công"),
        ("AI/ML correctness", "Đạt một phần", "Benchmark không leakage, có ablation và hold-out"),
        ("Software Engineering", "Đạt một phần", "Test/CI/Docker tái lập hoàn chỉnh"),
        ("Production / Business", "Chưa đạt", "Đo tải, bảo mật, privacy và observability"),
    ],
    [1800, 1800, 5760],
)
callout(
    doc,
    "Ưu tiên số 1",
    "Xây lại benchmark và đồng bộ số liệu. Nếu dữ liệu đánh giá không đáng tin, mọi tối ưu HyDE, router hay reranker đều không có cơ sở.",
    RED,
)

doc.add_heading("2. Tầng 1 — Problem", level=1)
doc.add_heading("2.1 Khóa định nghĩa bài toán", level=2)
doc.add_paragraph("Bổ sung vào README một hợp đồng đầu vào–đầu ra ngắn và có thể kiểm chứng:")
add_table(
    doc,
    ["Thành phần", "Định nghĩa đề xuất"],
    [
        ("Input", "Mô tả phim tự nhiên; tùy chọn thể loại và năm/khoảng năm"),
        ("Output", "Top-k phim, metadata, thứ hạng và điểm tương quan tương đối"),
        ("Người dùng", "Người nhớ nội dung/chủ đề nhưng quên tên; người khám phá phim"),
        ("Thành công", "Phim đúng trong top 5; latency và error rate đạt ngưỡng công bố"),
        ("Không cam kết", "final_score không phải xác suất phim đúng"),
    ],
    [2700, 6660],
)
doc.add_heading("2.2 Phân nhóm query thực tế", level=2)
add_table(
    doc,
    ["Nhóm", "Ví dụ", "Kỳ vọng"],
    [
        ("Plot memory", "Người cha liên lạc với con gái qua kệ sách", "Dense/Hybrid"),
        ("Exact entity", "Phim Nolan có Matthew McConaughey", "BM25"),
        ("Theme", "Phim về cô độc và mất mát", "Dense"),
        ("Constraint", "Phim chiến tranh khoảng năm 2005", "Filter + Hybrid"),
        ("Noisy", "phim robot danh nhau ngoai hanh tinh", "Normalization + Dense"),
        ("Multilingual", "Mô tả tiếng Việt", "Multilingual embedding hoặc đánh giá giới hạn"),
    ],
    [1600, 4560, 3200],
)
doc.add_heading("2.3 Cải thiện demo", level=2)
checklist(doc, [
    "Thêm GIF/video hoặc ảnh chụp luồng tìm kiếm thực tế.",
    "Cho thấy cùng một query qua BM25, Dense và Hybrid.",
    "Công bố cả ví dụ thành công lẫn ví dụ thất bại.",
    "Đổi nhãn final_score thành 'điểm tương quan tương đối'.",
    "Bổ sung thông báo khi kết quả yếu hoặc không chắc chắn.",
])

doc.add_heading("3. Tầng 2 — AI/ML correctness", level=1)
doc.add_heading("3.1 Xây lại evaluation dataset", level=2)
doc.add_paragraph(
    "Tập hiện tại chứa nhiều query gần như chép overview, tạo lợi thế bất thường cho BM25. Cần tạo query độc lập, "
    "không sao chép câu hoặc keyword trực tiếp từ document đích."
)
add_table(
    doc,
    ["Tập", "Kích thước gợi ý", "Mục đích"],
    [
        ("Development", "100 query", "Phân tích lỗi và thử thuật toán"),
        ("Validation", "100 query", "Chọn threshold và hyperparameter"),
        ("Hold-out test", "200 query", "Báo cáo cuối; không dùng để chỉnh cấu hình"),
    ],
    [2000, 2200, 5160],
)
doc.add_paragraph("Schema đề xuất cho mỗi query:")
callout(
    doc,
    "CSV schema",
    "query_id, query, query_type, relevant_movie_ids, language, difficulty, source",
)
checklist(doc, [
    "Không dùng nguyên văn overview hoặc câu đầu overview.",
    "Không đưa title vào query plot-memory.",
    "Không dùng một template lặp lại cho toàn bộ tập.",
    "Gắn source=human hoặc source=synthetic cho từng query.",
    "Cho phép nhiều phim liên quan nếu use case mang tính khám phá.",
])

doc.add_heading("3.2 Thiết lập baseline và ablation", level=2)
add_table(
    doc,
    ["Mã", "Pipeline", "Mục đích"],
    [
        ("B0", "Exact/title search", "Mốc tối thiểu"),
        ("B1", "BM25 only", "Baseline lexical"),
        ("B2", "Dense only", "Baseline semantic"),
        ("B3", "Hybrid RRF", "Giá trị của fusion"),
        ("B4", "Hybrid + reranker", "Giá trị reranking"),
        ("B5", "Hybrid + HyDE", "Giá trị query expansion"),
        ("B6", "Hybrid + HyDE + reranker", "Pipeline nâng cao"),
        ("B7", "Full + router", "Đánh đổi chất lượng/latency/cost"),
    ],
    [900, 3800, 4660],
)
doc.add_paragraph("Mỗi pipeline phải báo cáo cùng một tập và cùng cấu hình index:")
bullet(doc, "Hit@1, Hit@5, Hit@10 và MRR@10.")
bullet(doc, "Recall@50 trước rerank; nDCG@10 nếu có nhiều nhãn liên quan.")
bullet(doc, "Latency p50/p95/p99; tách encoding, Qdrant, HyDE và rerank.")
bullet(doc, "Error/fallback rate và chi phí LLM trung bình/query.")

doc.add_heading("3.3 Kiểm chứng HyDE", level=2)
doc.add_paragraph(
    "Chạy A/B giữa Hybrid + reranker và Hybrid + HyDE + reranker, sau đó tách theo query type. HyDE chỉ nên "
    "ở pipeline chính khi cải thiện đáng kể chất lượng hoặc giải quyết một nhóm query cụ thể."
)
callout(doc, "Quy tắc quyết định", "Nếu HyDE không tăng Hit@5/MRR rõ ràng nhưng tăng latency/cost, chuyển HyDE thành experiment tùy chọn hoặc loại bỏ.", RED)

doc.add_heading("3.4 Hiệu chỉnh Adaptive Router", level=2)
doc.add_paragraph("Ghi log validation theo schema:")
callout(doc, "Router dataset", "query_id, top_score, score_gap, route, correct_at_1, latency_ms, llm_called")
doc.add_paragraph("Grid search gợi ý:")
add_table(
    doc,
    ["Tham số", "Giá trị thử"],
    [
        ("minimum_score", "0.015, 0.020, 0.025, 0.030"),
        ("confidence_gap", "0.001, 0.003, 0.005, 0.010"),
    ],
    [3000, 6360],
)
doc.add_paragraph("Tiêu chí giữ router:")
checklist(doc, [
    "EASY precision ≥ 95%.",
    "EASY coverage ≥ 30%.",
    "Giảm p50 latency ≥ 20%.",
    "Hit@5 toàn hệ thống giảm không quá 1 điểm phần trăm.",
])
callout(doc, "Hiện trạng cần cảnh giác", "Artifact cũ ghi EASY chỉ 4/200 query (2%). Ở mức này router gần như không tạo giá trị tiết kiệm.", RED)

doc.add_heading("3.5 Làm rõ score và metric", level=2)
doc.add_paragraph(
    "Min-Max score chỉ thể hiện vị trí tương đối trong danh sách hiện tại. Khi chỉ có một kết quả, score bằng 1.0 "
    "không có nghĩa kết quả chắc chắn đúng."
)
bullet(doc, "API nên ưu tiên rank và score thô có tên rõ ràng: rrf_score, rerank_score.")
bullet(doc, "UI không gọi score là confidence hoặc xác suất.")
bullet(doc, "README phải ghi điểm không so sánh được giữa hai query khác nhau.")

doc.add_page_break()
doc.add_heading("4. Tầng 3 — Software Engineering", level=1)
doc.add_heading("4.1 Ma trận test cần bổ sung", level=2)
add_table(
    doc,
    ["Module", "Test bắt buộc"],
    [
        ("Query", "Unicode tiếng Việt; chỉ ký tự đặc biệt; query dài; whitespace"),
        ("Filter", "Năm đảo; biên 1888/2100; genre All; format sai"),
        ("Ranking", "Input rỗng; một nhánh rỗng; duplicate; score bằng nhau; top_n=0"),
        ("Search", "EASY; HARD; HyDE fallback; empty result; reranker failure"),
        ("Store", "Dense lỗi; sparse lỗi; cả hai lỗi; timeout"),
        ("Cache", "Hit; miss; eviction; deepcopy; concurrency"),
        ("API/UI", "422/503/504; XSS payload; payload thiếu trường"),
        ("Pipeline", "CSV rỗng; thiếu cột; movie_id lỗi; duplicate"),
    ],
    [2100, 7260],
)

doc.add_heading("4.2 Integration test", level=2)
number(doc, "Khởi động Qdrant local trong CI hoặc test fixture.")
number(doc, "Index một tập nhỏ khoảng 10–20 phim.")
number(doc, "Gọi pipeline search thật, không mock Qdrant.")
number(doc, "Kiểm tra phim kỳ vọng nằm trong top 5.")
number(doc, "Dọn collection sau test để bảo đảm idempotency.")

doc.add_heading("4.3 Khởi tạo FastAPI và health check", level=2)
bullet(doc, "Khởi tạo SearchService bằng FastAPI lifespan hoặc lazy singleton có Lock.")
bullet(doc, "Tách /live cho process và /ready cho Qdrant/model/collection.")
bullet(doc, "Không tải model nặng trước khi validation request hoàn thành.")
bullet(doc, "Readiness trả 503 khi index hoặc model chưa sẵn sàng.")

doc.add_heading("4.4 Error handling và timeout", level=2)
add_table(
    doc,
    ["Tình huống", "HTTP", "Hành vi"],
    [
        ("Input/filter sai", "422", "Thông báo cụ thể"),
        ("Qdrant/model lỗi", "503", "Không lộ stack trace"),
        ("Timeout tổng", "504", "Có request_id"),
        ("Lỗi không xác định", "500", "Log nội bộ, response an toàn"),
    ],
    [3000, 1200, 5160],
)
bullet(doc, "Đặt timeout Qdrant, Groq và tổng request.")
bullet(doc, "Fallback HyDE/reranker phải được ghi log và metric, không âm thầm che lỗi.")

doc.add_heading("4.5 Cấu hình, dependency và CI", level=2)
checklist(doc, [
    "Đưa collection/model/retrieval_k/candidate_k/threshold/cache size vào Settings.",
    "Chọn pyproject.toml làm nguồn dependency chính.",
    "Docker cài từ project metadata hoặc sinh requirements tự động.",
    "CI chạy Ruff, pytest và coverage; không chỉ lint.",
    "Thêm Qdrant service cho integration job.",
    "Thêm pip-audit hoặc công cụ dependency scanning.",
])

doc.add_heading("4.6 Docker và khả năng clone-run", level=2)
callout(doc, "Docker networking", "Trong container, QDRANT_URL phải là http://qdrant:6333, không phải localhost:6333.", RED)
checklist(doc, [
    "Thêm healthcheck Qdrant và depends_on.condition=service_healthy.",
    "Tạo service/profile indexer để lập chỉ mục trước khi search.",
    "Cung cấp dataset mẫu nhỏ hoặc script sinh fixture.",
    "Bổ sung LICENSE hoặc bỏ badge/link MIT.",
    "Xóa đường dẫn máy cá nhân trong báo cáo evaluation.",
    "Ghi rõ model tải ở lần chạy đầu và yêu cầu RAM/disk.",
])

doc.add_heading("5. Tầng 4 — Production / Business value", level=1)
doc.add_heading("5.1 Thiết kế phục vụ nhiều người dùng", level=2)
doc.add_paragraph("Đơn giản hóa kiến trúc runtime để tránh tải model hai lần:")
callout(doc, "Kiến trúc đề xuất", "Streamlit → FastAPI → SearchService → Qdrant; UI không gọi SearchService trực tiếp.")
bullet(doc, "FastAPI là nơi duy nhất giữ encoder/reranker trong memory.")
bullet(doc, "Dùng nhiều worker chỉ sau khi đo RAM và xác định model có thể nhân bản.")
bullet(doc, "Thêm concurrency limit/backpressure cho HARD route.")
bullet(doc, "Cache nhiều worker dùng Redis; prototype có thể dùng OrderedDict + Lock.")

doc.add_heading("5.2 Load test 100 users", level=2)
doc.add_paragraph("Trước hết định nghĩa tải: concurrent users, requests/second hoặc requests/minute. Sau đó chạy:")
add_table(
    doc,
    ["Concurrency", "Đo lường"],
    [
        ("1, 5, 10", "Baseline và warm/cold start"),
        ("25, 50", "Điểm bắt đầu queue và tăng p95"),
        ("100", "Throughput, p99, error rate, RAM/CPU/GPU"),
    ],
    [2400, 6960],
)
bullet(doc, "Kịch bản tải nên có EASY, HARD, input lỗi và cache hit.")
bullet(doc, "Đo riêng qdrant_ms, hyde_ms, rerank_ms và total_ms.")
bullet(doc, "Không tuyên bố chịu được 100 users trước khi có artifact load-test tái lập.")

doc.add_heading("5.3 Security và privacy", level=2)
add_table(
    doc,
    ["Rủi ro", "Cải thiện"],
    [
        ("API công khai", "Authentication/API key và rate limiting"),
        ("HARD route tốn chi phí", "Quota riêng cho LLM route"),
        ("Query gửi Groq", "Disclosure rõ ràng và ENABLE_HYDE=false"),
        ("Qdrant public", "Private network; không expose 6333 nếu không cần"),
        ("Log nhạy cảm", "Không log nguyên query; dùng request_id/hash khi phù hợp"),
        ("Dependency", "pip-audit/Dependabot và cập nhật định kỳ"),
    ],
    [3300, 6060],
)

doc.add_heading("5.4 Observability", level=2)
doc.add_paragraph("Mỗi request nên có structured log tối thiểu:")
callout(doc, "Log fields", "request_id, route, candidate_count, cache_hit, qdrant_ms, hyde_ms, rerank_ms, total_ms, fallback, status")
doc.add_paragraph("Metrics gợi ý:")
bullet(doc, "search_requests_total, search_errors_total, search_latency_seconds.")
bullet(doc, "search_route_total, hyde_fallback_total, qdrant_errors_total.")
bullet(doc, "cache_hit_total và zero_result_total.")

doc.add_heading("5.5 Business metrics", level=2)
bullet(doc, "Tỷ lệ người dùng chọn phim trong top 1/top 5.")
bullet(doc, "Search success rate và zero-result rate.")
bullet(doc, "Query refinement rate và time-to-first-useful-result.")
bullet(doc, "Chi phí cho mỗi search thành công.")
bullet(doc, "Feedback đúng/sai gắn với request_id, không lưu dữ liệu cá nhân thừa.")

doc.add_page_break()
doc.add_heading("6. Lộ trình triển khai theo phase", level=1)
add_table(
    doc,
    ["Phase", "Thời lượng gợi ý", "Đầu ra bắt buộc"],
    [
        ("1. Correctness", "3–5 ngày", "Dataset mới; split; metric đồng bộ; limitation"),
        ("2. AI evidence", "3–5 ngày", "Baseline, ablation, router calibration"),
        ("3. Reliability", "3–5 ngày", "Unit/integration test; health; timeout; CI"),
        ("4. Reproducibility", "2–4 ngày", "Docker clone-to-run; sample data; license"),
        ("5. Production", "5–10 ngày", "Load test; rate limit; privacy; monitoring"),
    ],
    [1700, 2100, 5560],
)
doc.add_heading("Phase 1 — Correctness", level=2)
checklist(doc, [
    "Tạo development/validation/test không copy overview.",
    "Chạy lại BM25 và Dense baseline.",
    "Đồng bộ README với output script.",
    "Ghi rõ final_score và limitations.",
])
doc.add_heading("Phase 2 — AI evidence", level=2)
checklist(doc, [
    "Chạy B0–B7 trên cùng hold-out.",
    "Tách kết quả theo query type/language/difficulty.",
    "Quyết định giữ hoặc bỏ HyDE/router bằng số liệu.",
])
doc.add_heading("Phase 3 — Reliability", level=2)
checklist(doc, [
    "Bổ sung HARD/failure/cache tests.",
    "Integration test Qdrant.",
    "CI chạy test và coverage.",
    "Health/readiness và timeout.",
])
doc.add_heading("Phase 4 — Reproducibility", level=2)
checklist(doc, [
    "Docker networking và indexer.",
    "Dataset mẫu và smoke test.",
    "Clean-install từ clone mới.",
])
doc.add_heading("Phase 5 — Production", level=2)
checklist(doc, [
    "UI gọi API; đo tải 1–100 concurrency.",
    "Authentication/rate limiting/privacy disclosure.",
    "Structured logging và metrics.",
])

doc.add_heading("7. Definition of Done và checklist nghiệm thu", level=1)
doc.add_heading("7.1 Portfolio-ready", level=2)
checklist(doc, [
    "ruff check . đạt.",
    "Toàn bộ unit test và integration test đạt.",
    "CI chạy lint + test.",
    "Docker chạy từ clone sạch đến kết quả search.",
    "Hold-out benchmark không leakage.",
    "README metric khớp artifact.",
    "Có baseline, ablation, limitation và demo thất bại.",
    "Không có secret hoặc path máy cá nhân.",
])
doc.add_heading("7.2 Gần production-ready", level=2)
checklist(doc, [
    "Load test 100 concurrent users có p95/p99 và error rate.",
    "Authentication, rate limiting và timeout hoạt động.",
    "Readiness phản ánh Qdrant/model/index thật.",
    "Query gửi Groq có disclosure và opt-out.",
    "Monitoring, backup/index versioning và rollback được mô tả.",
])
callout(
    doc,
    "Nguyên tắc cuối",
    "Chỉ gọi một thành phần là cần thiết khi nó cải thiện metric, latency, cost hoặc reliability trên benchmark đáng tin. Nếu không, loại bỏ để hệ thống đơn giản hơn.",
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT.resolve())
